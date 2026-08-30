import io
import uuid
import json
from typing import Any, Dict, List, Optional, Tuple
from pypdf import PdfReader
from fastembed import TextEmbedding
from langchain_core.messages import SystemMessage
from src.config.llm_config import get_classifier_model
from src.repositories.document_repository import (
    insert_document,
    insert_document_chunks,
    list_user_documents,
    delete_user_document,
    delete_documents_by_filename,
    delete_all_documents,
    similarity_search_chunks,
)

_embedding_model: TextEmbedding | None = None

def get_embedding_model() -> TextEmbedding:
    global _embedding_model
    if _embedding_model is None:
        _embedding_model = TextEmbedding(model_name="BAAI/bge-small-en-v1.5")
    return _embedding_model

_CLASSIFIER_PROMPT = """You are a strict document relevance validator for an executive AI scheduling assistant (Cadence).
Cadence ONLY ingests work and scheduling-related documents to help schedule meetings, understand team availability, project deadlines, sprint plans, conference schedules, or meeting agendas.

REJECT: Resumes, CVs, invoices, bank receipts, general fiction/novels, textbooks, code dumps, personal medical logs, or unrelated files.
ACCEPT: Meeting agendas, meeting minutes/notes, project sprint plans, team calendars, event itineraries, scheduling briefs.

Analyze the document sample:
\"\"\"{sample_text}\"\"\"

Reply strictly in this JSON format and nothing else:
{{
  "is_relevant": true/false,
  "doc_type": "Agenda / Sprint Plan / Meeting Notes / Irrelevant",
  "summary": "1-2 sentence summary of what this document covers",
  "rejection_reason": "Reason if is_relevant is false, otherwise empty string"
}}"""


# ---------------------------------------------------------------------------
# Text extractors
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text)
    return "\n\n".join(pages_text).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file including table content."""
    try:
        from docx import Document  # type: ignore[import]
        doc = Document(io.BytesIO(file_bytes))
        paragraphs: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Could not read the DOCX file: {e}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract plain text from a TXT file, trying common encodings."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(
        "Could not decode the text file. Please ensure it uses UTF-8 encoding."
    )


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 80) -> List[str]:
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks if chunks else [text]


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

async def validate_document_relevance(text_sample: str) -> Dict[str, Any]:
    model = get_classifier_model()
    sample = text_sample[:2500]
    response = await model.ainvoke(
        [SystemMessage(content=_CLASSIFIER_PROMPT.format(sample_text=sample))]
    )
    content = str(response.content).strip()

    if content.startswith("```"):
        content = content.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
    try:
        return json.loads(content)
    except Exception:
        is_relevant = "true" in content.lower()
        return {
            "is_relevant": is_relevant,
            "doc_type": "Work Document" if is_relevant else "Irrelevant",
            "summary": "Processed work document",
            "rejection_reason": (
                "" if is_relevant
                else "Document not recognized as a valid scheduling/work document."
            ),
        }


# ---------------------------------------------------------------------------
# Core ingestion — supports PDF, DOCX, TXT
# ---------------------------------------------------------------------------

async def process_and_ingest_document(
    oauth_user_id: str,
    filename: str,
    file_bytes: bytes,
) -> Tuple[bool, str, Dict[str, Any]]:
    """Validate, extract, classify, chunk, and index a document.
    Supports PDF, DOCX, and TXT. Returns (success, message, doc_metadata)."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Extract text based on file type
    try:
        if ext == "pdf":
            raw_text = extract_text_from_pdf(file_bytes)
        elif ext == "docx":
            raw_text = extract_text_from_docx(file_bytes)
        elif ext == "txt":
            raw_text = extract_text_from_txt(file_bytes)
        else:
            return (
                False,
                f"Unsupported file type '.{ext}'. Please upload a PDF, DOCX, or TXT file.",
                {},
            )
    except ValueError as e:
        return False, str(e), {}

    if not raw_text or len(raw_text.strip()) < 50:
        return False, "The document appears to be empty or contains no readable text.", {}

    # Classify relevance
    validation = await validate_document_relevance(raw_text)
    if not validation.get("is_relevant", False):
        reason = (
            validation.get("rejection_reason")
            or "Document is not a work scheduling or agenda document."
        )
        return False, f"Rejection: {reason}", validation

    # Deduplicate: remove any prior version of same filename
    await delete_documents_by_filename(oauth_user_id, filename)

    doc_id = str(uuid.uuid4())
    chunks_text = chunk_text(raw_text)

    embedder = get_embedding_model()
    embeddings_gen = list(embedder.embed(chunks_text))

    chunks_data = [
        {"index": idx, "content": chunk_str, "embedding": embeddings_gen[idx].tolist()}
        for idx, chunk_str in enumerate(chunks_text)
    ]

    # Use exact same insert_document signature as the original
    await insert_document(
        doc_id=doc_id,
        oauth_user_id=oauth_user_id,
        filename=filename,
        file_size=len(file_bytes),
        doc_type=validation.get("doc_type", "Work Document"),
        summary=validation.get("summary", "Uploaded work document."),
    )
    await insert_document_chunks(doc_id=doc_id, oauth_user_id=oauth_user_id, chunks=chunks_data)

    return True, "Document validated and indexed successfully.", {
        "id": doc_id,
        "filename": filename,
        "doc_type": validation.get("doc_type"),
        "summary": validation.get("summary"),
        "chunks_count": len(chunks_text),
    }


# Keep old name as alias so document_routes.py doesn't need changing
async def process_and_ingest_pdf(
    oauth_user_id: str,
    filename: str,
    file_bytes: bytes,
) -> Tuple[bool, str, Dict[str, Any]]:
    return await process_and_ingest_document(oauth_user_id, filename, file_bytes)


# ---------------------------------------------------------------------------
# Search, list, delete
# ---------------------------------------------------------------------------

async def search_rag_context(
    oauth_user_id: str,
    query: str,
    top_k: int = 3,
    filename: Optional[str] = None,
) -> List[Dict[str, Any]]:
    embedder = get_embedding_model()
    query_embedding = list(embedder.embed([query]))[0].tolist()
    return await similarity_search_chunks(
        oauth_user_id, query_embedding, top_k=top_k, filename=filename
    )


async def get_user_docs(oauth_user_id: str) -> List[Dict[str, Any]]:
    return await list_user_documents(oauth_user_id)


async def remove_user_doc(doc_id: str, oauth_user_id: str) -> bool:
    return await delete_user_document(doc_id, oauth_user_id)


async def remove_documents_by_name(oauth_user_id: str, filename: str) -> None:
    await delete_documents_by_filename(oauth_user_id, filename)


async def remove_all_documents(oauth_user_id: str) -> int:
    return await delete_all_documents(oauth_user_id)